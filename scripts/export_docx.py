#!/usr/bin/env python3
"""Generate a DOCX from literature search results (JSON stdin -> DOCX file)."""
import sys, json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shade = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shade)

def build_docx(data, output_path):
    query = data.get("query", "Literature")
    results = data.get("results", [])

    doc = Document()
    doc.core_properties.author = "Perplexity Computer"
    doc.core_properties.title = f"Literature Search: {query}"

    # Page setup - landscape
    for section in doc.sections:
        section.orientation = 1  # landscape
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Title
    title_p = doc.add_paragraph()
    run = title_p.add_run(f"文獻搜尋結果：{query}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x52, 0x76)

    # Subtitle
    sub_p = doc.add_paragraph()
    run = sub_p.add_run(
        f"共 {len(results)} 篇 · IF ≥ {data.get('minIF', 4)} · 資料來源：OpenAlex")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if not results:
        doc.add_paragraph("未找到符合條件的文獻")
        doc.save(output_path)
        return

    # Table
    headers = ["#", "標題", "作者", "期刊", "IF", "日期", "連結"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1a5276")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for idx, r in enumerate(results, 1):
        row = table.add_row()
        cells = row.cells

        # #
        cells[0].text = str(idx)
        cells[0].paragraphs[0].runs[0].font.size = Pt(7) if cells[0].paragraphs[0].runs else None

        # Title + abstract snippet
        title_cell = cells[1]
        title_cell.text = ""
        tp = title_cell.paragraphs[0]
        tr = tp.add_run(r.get("title", ""))
        tr.bold = True
        tr.font.size = Pt(8)
        abstract = (r.get("abstract", "") or "")[:150]
        if abstract:
            tp2 = title_cell.add_paragraph()
            ar = tp2.add_run(abstract + "...")
            ar.font.size = Pt(6.5)
            ar.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Author
        cells[2].text = ""
        ap = cells[2].paragraphs[0]
        ar2 = ap.add_run((r.get("author", "") or "")[:120])
        ar2.font.size = Pt(7)

        # Journal
        cells[3].text = ""
        jp = cells[3].paragraphs[0]
        jr = jp.add_run(r.get("journal", ""))
        jr.font.size = Pt(7)

        # IF
        cells[4].text = ""
        ip = cells[4].paragraphs[0]
        ir2 = ip.add_run(str(r.get("impactFactor", "")))
        ir2.font.size = Pt(7)
        ir2.bold = True
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        cells[5].text = ""
        dp = cells[5].paragraphs[0]
        dr2 = dp.add_run((r.get("publicationDate", "") or "")[:10])
        dr2.font.size = Pt(7)

        # Link (hyperlink)
        pub_url = r.get("pubUrl", "")
        cells[6].text = ""
        if pub_url:
            lp = cells[6].paragraphs[0]
            add_hyperlink(lp, pub_url, "Link")

        # Alternate row shading
        if idx % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "f8f9fa")

        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Set column widths approximately
    widths = [Cm(1), Cm(8), Cm(5.5), Cm(3.5), Cm(1.2), Cm(2), Cm(5.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.save(output_path)


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:color w:val="1a73e8"/>'
        f'      <w:u w:val="single"/>'
        f'      <w:sz w:val="14"/>'
        f'    </w:rPr>'
        f'    <w:t>{text}</w:t>'
        f'  </w:r>'
        f'</w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


if __name__ == "__main__":
    data = json.load(sys.stdin)
    output = data.get("outputPath", "/tmp/export.docx")
    build_docx(data, output)
    print(json.dumps({"ok": True, "path": output}))
